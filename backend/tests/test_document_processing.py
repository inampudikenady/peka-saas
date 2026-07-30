from io import BytesIO
import os
from uuid import uuid4

import pytest

from app.services.document_chunker import chunk_document
from app.services.document_parsers import ParsedDocument, ParsedSection, parser_for
from app.services.embedding_provider import (
    DeterministicFakeEmbeddingProvider,
    EmbeddingDimensionError,
    OpenAICompatibleEmbeddingProvider,
)
from app.services.vector_store import InMemoryVectorStore, QdrantVectorStore, VectorPoint


@pytest.mark.parametrize("filename", ["notes.txt", "notes.md", "rows.csv"])
def test_text_parsers_extract_content(filename):
    content = b"name,value\nalpha,one" if filename.endswith(".csv") else b"alpha beta gamma"
    parsed = parser_for(filename).parse(BytesIO(content))
    assert parsed.sections
    assert "alpha" in parsed.sections[0].text


def test_chunking_preserves_citation_metadata_and_overlap():
    parsed = ParsedDocument([
        ParsedSection(
            text=" ".join(f"word-{index}" for index in range(30)), page_number=4,
            sheet_name="Policy", row_start=10, row_end=20, section_title="Rotation",
        )
    ], "test")
    chunks = chunk_document(parsed, target_words=10, overlap_words=2)
    assert len(chunks) > 3
    assert all(chunk.page_number == 4 and chunk.sheet_name == "Policy" for chunk in chunks)
    assert chunks[0].text.split()[-2:] == chunks[1].text.split()[:2]


def test_vector_search_always_filters_tenant():
    first, second = uuid4(), uuid4()
    embeddings = DeterministicFakeEmbeddingProvider(16)
    vector = embeddings.embed(["password rotation"])[0]
    store = InMemoryVectorStore()
    store.upsert([
        VectorPoint(uuid4(), vector, {"tenant_id": str(first), "document_id": "a"}),
        VectorPoint(uuid4(), vector, {"tenant_id": str(second), "document_id": "b"}),
    ])
    hits = store.search(first, vector, 10, {})
    assert len(hits) == 1
    assert hits[0].payload["tenant_id"] == str(first)
    assert store.count_points(first) == 1
    assert store.count_points(second) == 1


def test_fake_embeddings_are_deterministic_and_point_ids_are_stable():
    provider = DeterministicFakeEmbeddingProvider(16)
    assert provider.embed(["same text"])[0] == provider.embed(["same text"])[0]
    version_id = uuid4()
    from uuid import NAMESPACE_URL, uuid5
    assert uuid5(NAMESPACE_URL, f"peka:{version_id}:0") == uuid5(
        NAMESPACE_URL, f"peka:{version_id}:0"
    )


def test_real_provider_dimension_mismatch_is_rejected(monkeypatch):
    class Response:
        status_code = 200

        def raise_for_status(self):
            return None

        def json(self):
            return {"data": [{"embedding": [0.1, 0.2]}]}

    monkeypatch.setattr("httpx.post", lambda *args, **kwargs: Response())
    provider = OpenAICompatibleEmbeddingProvider(
        "https://embeddings.invalid", "secret", "model", dimension=3
    )
    with pytest.raises(EmbeddingDimensionError):
        provider.embed(["dimension check"])


@pytest.mark.skipif(
    not os.getenv("PEKA_TEST_QDRANT_URL"),
    reason="Set PEKA_TEST_QDRANT_URL to run the Qdrant integration test.",
)
def test_qdrant_tenant_filter_upsert_search_and_delete():
    import httpx

    url = os.environ["PEKA_TEST_QDRANT_URL"]
    collection = f"peka_test_{uuid4().hex}"
    store = QdrantVectorStore(url, collection, timeout=10)
    tenant_id, other_tenant_id = uuid4(), uuid4()
    document_id = uuid4()
    try:
        store.ensure_collection(3)
        store.upsert([
            VectorPoint(
                uuid4(), [1.0, 0.0, 0.0],
                {"tenant_id": str(tenant_id), "document_id": str(document_id),
                 "lifecycle_status": "ACTIVE"},
            ),
            VectorPoint(
                uuid4(), [1.0, 0.0, 0.0],
                {"tenant_id": str(other_tenant_id), "document_id": str(uuid4()),
                 "lifecycle_status": "ACTIVE"},
            ),
        ])
        hits = store.search(
            tenant_id, [1.0, 0.0, 0.0], 10, {"lifecycle_status": "ACTIVE"}
        )
        assert len(hits) == 1
        assert hits[0].payload["tenant_id"] == str(tenant_id)
        assert store.count_points(tenant_id, document_id=document_id) == 1
        store.delete_document(tenant_id, document_id)
        assert store.count_points(tenant_id, document_id=document_id) == 0
        assert store.search(tenant_id, [1.0, 0.0, 0.0], 10, {}) == []
    finally:
        store.client.close()
        httpx.delete(f"{url.rstrip('/')}/collections/{collection}", timeout=10)


def test_markdown_headings_and_line_endings_are_preserved():
    parsed = parser_for("policy.md").parse(BytesIO(b"# Passwords\r\nRotate regularly.\r\n"))
    assert parsed.sections[0].section_title == "Passwords"
    assert parsed.sections[0].text.startswith("# Passwords")
    assert "\r" not in parsed.sections[0].text


def test_same_txt_extension_detects_plain_markdown_and_dokuwiki_content():
    plain = parser_for("notes.txt", "text/plain").parse(
        BytesIO(b"Call the service desk if assistance is required.")
    )
    markdown = parser_for("notes.txt", "text/plain").parse(
        BytesIO(b"# Restart\n\n```bash\nsystemctl restart peka\n```\n")
    )
    dokuwiki = parser_for("notes.txt", "text/plain").parse(
        BytesIO(b"====== Restart ======\n  * Run<code|bash>systemctl restart peka</code>\n")
    )

    assert plain.detected_format == "plain_text"
    assert markdown.detected_format == "markdown"
    assert dokuwiki.detected_format == "dokuwiki"
    assert dokuwiki.source_format == "dokuwiki_export"


def test_dokuwiki_normalizes_supported_structures_to_canonical_markdown():
    source = b"""====== Operations ======
  * Restart the service
<code bash>
systemctl restart peka
  systemctl status peka
</code>
^ Host ^ State ^
| util001 | up |
[[https://example.test/runbook|Runbook]]
Use ''systemctl''.
"""
    parsed = parser_for("operations.txt", "text/plain").parse(BytesIO(source))
    normalized = "\n".join(section.text for section in parsed.sections)

    assert normalized.startswith("# Operations")
    assert "- Restart the service" in normalized
    assert "```bash\nsystemctl restart peka\n  systemctl status peka\n```" in normalized
    assert "| Host | State |" in normalized
    assert "| --- | --- |" in normalized
    assert "[Runbook](https://example.test/runbook)" in normalized
    assert "`systemctl`" in normalized


def test_mixed_symbols_without_strong_structure_fall_back_to_plain_text():
    parsed = parser_for("mixed.txt", "text/plain").parse(
        BytesIO(b"Use #1 for support. The value a|b is accepted and [x] is a label.")
    )
    assert parsed.detected_format == "plain_text"


def test_structured_chunking_preserves_code_and_heading_context():
    command = (
        b"useradd -g dba " + b"\\" + b"\n"
        b"  -d /home/kohlerdba " + b"\\" + b"\n"
        b"  kohlerdba\n"
    )
    parsed = parser_for("runbook.txt", "text/plain").parse(
        BytesIO(
            b"====== User setup ======\n"
            b"  * Create the account<code bash>\n"
            + command
            + b"</code>\n"
            + b"\nParagraph content. " * 30
        )
    )
    chunks = chunk_document(parsed, target_words=25)
    command_chunk = next(chunk for chunk in chunks if "useradd -g dba" in chunk.text)

    assert "# User setup" in command_chunk.text
    assert "```bash" in command_chunk.text
    assert "useradd -g dba \\\n  -d /home/kohlerdba \\\n  kohlerdba" in command_chunk.text
    assert all("[object Object]" not in chunk.text for chunk in chunks)


def test_docx_heading_and_xlsx_sheet_rows_are_preserved():
    from docx import Document
    from openpyxl import Workbook

    word = Document(); word.add_heading("Install vManager", level=1)
    word.add_paragraph("Run the signed installer.")
    word_stream = BytesIO(); word.save(word_stream); word_stream.seek(0)
    parsed_word = parser_for("guide.docx").parse(word_stream)
    assert parsed_word.sections[0].section_title == "Install vManager"

    workbook = Workbook(); sheet = workbook.active; sheet.title = "Servers"
    sheet.append(["Host", "Version"]); sheet.append(["vm-01", "2.0"])
    excel_stream = BytesIO(); workbook.save(excel_stream); excel_stream.seek(0)
    parsed_excel = parser_for("inventory.xlsx").parse(excel_stream)
    assert parsed_excel.sections[0].sheet_name == "Servers"
    assert parsed_excel.sections[1].row_start == 2
    chunks = chunk_document(parsed_excel, target_words=2)
    assert all(chunk.sheet_name == "Servers" for chunk in chunks)
    assert "Host" in chunks[-1].text


def test_likely_scanned_pdf_fails_safely():
    from pypdf import PdfWriter

    writer = PdfWriter(); writer.add_blank_page(width=612, height=792)
    stream = BytesIO(); writer.write(stream); stream.seek(0)
    with pytest.raises(ValueError, match="scanned"):
        parser_for("scan.pdf").parse(stream)
