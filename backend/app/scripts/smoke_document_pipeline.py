"""Development-only live smoke test for the document vertical slice."""

from __future__ import annotations

import hashlib
import io
import json
import time
from datetime import UTC, datetime
from uuid import uuid4

import httpx
from docx import Document as WordDocument
from openpyxl import Workbook
from sqlalchemy import select

from app.core.config import settings
from app.core.security import create_tenant_access_token
from app.db.session import SessionLocal
from app.models.connector import ConnectorRegistrationToken, ManagedConnector
from app.models.document import Document, DocumentVersion
from app.models.tenant import Tenant
from app.models.tenant_user import TenantUser, TenantUserRole
from app.repositories.connector_repository import ConnectorRepository
from app.repositories.tenant_repository import TenantRepository
from app.services.connector_service import ConnectorService
from app.services.provider_factory import object_storage, vector_store


def _pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(escaped) + 35} >>\nstream\nBT /F1 12 Tf 72 720 Td ({escaped}) Tj ET\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = io.BytesIO(); output.write(b"%PDF-1.4\n")
    offsets = [0]
    for number, value in enumerate(objects, start=1):
        offsets.append(output.tell()); output.write(f"{number} 0 obj\n{value}\nendobj\n".encode())
    xref = output.tell(); output.write(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]: output.write(f"{offset:010d} 00000 n \n".encode())
    output.write(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode())
    return output.getvalue()


def _fixtures() -> dict[str, tuple[bytes, str]]:
    word = WordDocument(); word.add_heading("Password Policy", level=1)
    word.add_paragraph("The service account password rotates every 90 days.")
    word_bytes = io.BytesIO(); word.save(word_bytes)
    workbook = Workbook(); sheet = workbook.active; sheet.title = "Rotation"
    sheet.append(["Account", "Days"]); sheet.append(["Linux service", 90])
    excel_bytes = io.BytesIO(); workbook.save(excel_bytes)
    return {
        "smoke.txt": (b"PEKA smoke marker: rotate the Linux service password every 90 days.", "text/plain"),
        "smoke.pdf": (_pdf("PEKA PDF smoke marker requires quarterly access review."), "application/pdf"),
        "smoke.docx": (word_bytes.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        "smoke.xlsx": (excel_bytes.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    }


def main() -> None:
    if settings.environment.lower() not in {"local", "development", "test"}:
        raise SystemExit("Refusing to run outside a development environment")
    base_url = "http://127.0.0.1:8000"
    session = SessionLocal(); connector_id = None; token_id = None
    document_ids: list = []; object_keys: list[str] = []
    try:
        tenant = session.scalar(select(Tenant).where(Tenant.slug == "vitwo"))
        if tenant is None: raise RuntimeError("Tenant 'vitwo' does not exist")
        admin = session.scalar(select(TenantUser).where(
            TenantUser.tenant_id == tenant.id, TenantUser.role == TenantUserRole.TENANT_ADMIN,
            TenantUser.is_active.is_(True),
        ))
        if admin is None: raise RuntimeError("VITWO has no active tenant administrator")
        connector_service = ConnectorService(ConnectorRepository(session), TenantRepository(session))
        registration = connector_service.create_registration_token(tenant.id, admin, None)
        token_id = registration.id; instance_id = uuid4()
        client = httpx.Client(base_url=base_url, timeout=30)
        registered = client.post("/api/v1/connectors/register", headers={"Host": "backend.internal"}, json={
            "registration_token": registration.registration_token, "connector_name": "Document Smoke Connector",
            "connector_version": "smoke-1", "environment": "development",
            "instance_id": str(instance_id), "capabilities": ["filesystem_documents"],
        }); registered.raise_for_status(); credentials = registered.json(); connector_id = credentials["connector_id"]
        connector_headers = {"Authorization": f"Bearer {credentials['connector_secret']}",
                             "X-PEKA-Connector-ID": connector_id, "Host": "127.0.0.1:8000"}
        for filename, (content, mime_type) in _fixtures().items():
            digest = f"sha256:{hashlib.sha256(content).hexdigest()}"
            metadata = {"source_id": "smoke-source", "document_key": filename,
                        "relative_path": f"smoke/{filename}", "filename": filename,
                        "mime_type": mime_type, "size_bytes": len(content), "content_hash": digest,
                        "modified_at": datetime.now(UTC).isoformat(), "operation": "upsert",
                        "connector_version": "smoke-1"}
            response = client.post(f"/api/v1/connectors/{connector_id}/documents",
                headers={**connector_headers, "Idempotency-Key": f"smoke-{uuid4()}"},
                files={"metadata": (None, json.dumps(metadata), "application/json"),
                       "file": (filename, content, mime_type)})
            response.raise_for_status(); acknowledgement = response.json()
            assert acknowledgement["content_hash"] == digest
            document_ids.append(acknowledgement["document_id"])
        deadline = time.monotonic() + 45
        states: dict[str, str] = {}
        while time.monotonic() < deadline:
            status_response = client.get(
                f"/api/v1/connectors/{connector_id}/documents/status", headers=connector_headers
            ); status_response.raise_for_status()
            states = {item["document_id"]: item["ingestion_status"] for item in status_response.json()}
            if len(states) == 4 and all(value == "INDEXED" for value in states.values()): break
            time.sleep(1)
        if not states or any(value != "INDEXED" for value in states.values()):
            raise RuntimeError(f"Documents did not index: {states}")
        tenant_token = create_tenant_access_token(admin.id, admin.username or admin.email, tenant.id)
        tenant_headers = {"Authorization": f"Bearer {tenant_token}"}
        search = client.post("/t/vitwo/api/v1/tenant/search", headers=tenant_headers,
                             json={"query": "Linux service password rotation", "top_k": 8, "filters": {}})
        search.raise_for_status(); results = search.json()["results"]
        if not results or any(item["document_id"] not in document_ids for item in results):
            raise RuntimeError("Tenant search returned unexpected results")
        print(json.dumps({"tenant_id": str(tenant.id), "connector_id": connector_id,
                          "formats_indexed": sorted(_fixtures()), "search_results": len(results),
                          "result": "passed"}))
    finally:
        session.rollback()
        if connector_id is not None:
            connector = session.get(ManagedConnector, connector_id)
            if connector is not None:
                documents = list(session.scalars(select(Document).where(Document.connector_id == connector.id)))
                vectors = vector_store(); storage = object_storage()
                for document in documents:
                    vectors.delete_document(document.tenant_id, document.id)
                    versions = session.scalars(select(DocumentVersion).where(DocumentVersion.document_id == document.id))
                    for version in versions: object_keys.append(version.object_key)
                session.delete(connector)
        if token_id is not None:
            token = session.get(ConnectorRegistrationToken, token_id)
            if token is not None: session.delete(token)
        session.commit()
        storage = object_storage()
        for key in object_keys: storage.delete(key)
        session.close()


if __name__ == "__main__":
    main()
