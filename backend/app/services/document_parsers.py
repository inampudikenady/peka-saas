"""Normalized parsers for the first supported document formats."""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, BinaryIO, Protocol

from app.services.structured_text import detect_text_format, normalize_dokuwiki


@dataclass(frozen=True)
class ParsedSection:
    text: str
    page_number: int | None = None
    sheet_name: str | None = None
    row_start: int | None = None
    row_end: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    sections: list[ParsedSection]
    parser_name: str
    parser_version: str = "1"
    title: str | None = None
    detected_format: str | None = None
    detection_confidence: float | None = None
    detection_reason: str | None = None
    source_format: str | None = None


class Parser(Protocol):
    def parse(self, stream: BinaryIO) -> ParsedDocument: ...


def _markdown_sections(content: str) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    heading: str | None = None
    buffer: list[str] = []
    in_fence = False
    for line in content.splitlines():
        if line.lstrip().startswith(("```", "~~~")):
            in_fence = not in_fence
        heading_match = (
            None if in_fence else re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line)
        )
        if heading_match:
            if buffer:
                sections.append(
                    ParsedSection("\n".join(buffer).strip(), section_title=heading)
                )
            heading = heading_match.group(2).strip() or None
            buffer = [line]
        else:
            buffer.append(line)
    if buffer or heading:
        sections.append(ParsedSection("\n".join(buffer).strip(), section_title=heading))
    return [section for section in sections if section.text.strip()]


class TextParser:
    def __init__(self, filename: str = "document.txt", mime_type: str | None = None) -> None:
        self.filename = filename
        self.mime_type = mime_type

    def parse(self, stream: BinaryIO) -> ParsedDocument:
        raw = stream.read()
        try:
            content = raw.decode("utf-8")
        except UnicodeDecodeError:
            content = raw.decode("utf-8", errors="replace")
        content = content.replace("\r\n", "\n").replace("\r", "\n")
        detection = detect_text_format(self.filename, self.mime_type, content)
        normalized = normalize_dokuwiki(content) if detection.detected_format == "dokuwiki" else content
        sections = (
            _markdown_sections(normalized)
            if detection.detected_format in {"markdown", "dokuwiki"}
            else [ParsedSection(text=normalized)]
        )
        return ParsedDocument(
            sections=sections,
            parser_name=detection.detected_format,
            parser_version="2",
            title=next((item.section_title for item in sections if item.section_title), None),
            detected_format=detection.detected_format,
            detection_confidence=detection.confidence,
            detection_reason=detection.reason,
            source_format=detection.source_format,
        )


class CsvParser:
    def parse(self, stream: BinaryIO) -> ParsedDocument:
        raw = stream.read()
        if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
            text = raw.decode("utf-16")
        else:
            try:
                text = raw.decode("utf-8-sig")
            except UnicodeDecodeError:
                text = raw.decode("cp1252", errors="replace")
        rows = list(csv.reader(io.StringIO(text, newline="")))
        header = rows[0] if rows else []
        sections: list[ParsedSection] = []
        for index, row in enumerate(rows[1:] if header else rows, start=2 if header else 1):
            values = [
                f"{header[column]}: {value}" if column < len(header) else value
                for column, value in enumerate(row)
            ]
            sections.append(ParsedSection(
                text="\n".join(values), row_start=index, row_end=index,
                metadata={"headers": header},
            ))
        return ParsedDocument(sections, "csv", "1")


class PdfParser:
    def parse(self, stream: BinaryIO) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(stream)
        if reader.is_encrypted:
            raise ValueError("Encrypted PDF documents are unsupported")
        sections = [
            ParsedSection(text=page.extract_text() or "", page_number=index)
            for index, page in enumerate(reader.pages, start=1)
        ]
        if sections and sum(len(section.text.strip()) for section in sections) < len(sections) * 10:
            raise ValueError("PDF appears scanned or contains insufficient extractable text")
        title = str(reader.metadata.title) if reader.metadata and reader.metadata.title else None
        return ParsedDocument(sections, "pypdf", "1", title)


class DocxParser:
    def parse(self, stream: BinaryIO) -> ParsedDocument:
        from docx import Document

        document = Document(stream)
        sections: list[ParsedSection] = []
        title: str | None = None
        buffer: list[str] = []
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading"):
                if buffer:
                    sections.append(ParsedSection(text="\n".join(buffer), section_title=title))
                    buffer = []
                title = paragraph.text.strip() or None
            elif paragraph.text.strip():
                buffer.append(paragraph.text)
        if buffer or title:
            sections.append(ParsedSection(text="\n".join(buffer), section_title=title))
        for table_index, table in enumerate(document.tables, start=1):
            rows = [" | ".join(cell.text for cell in row.cells) for row in table.rows]
            sections.append(
                ParsedSection(text="\n".join(rows), section_title=f"Table {table_index}")
            )
        return ParsedDocument(sections, "python-docx", "1", sections[0].section_title if sections else None)


class XlsxParser:
    def parse(self, stream: BinaryIO) -> ParsedDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(stream, read_only=True, data_only=True)
        sections: list[ParsedSection] = []
        for sheet in workbook.worksheets:
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    sections.append(
                        ParsedSection(
                            text=" | ".join(values),
                            sheet_name=sheet.title,
                            row_start=row_number,
                            row_end=row_number,
                        )
                    )
        return ParsedDocument(sections, "openpyxl", "1", workbook.properties.title)


class ParserRegistry:
    def __init__(self) -> None:
        self._parsers: dict[str, Parser] = {}

    def register(self, extension: str, parser: Parser) -> None:
        self._parsers[extension.lower()] = parser

    def get(self, filename: str, mime_type: str | None = None) -> Parser:
        extension = Path(filename).suffix.lower()
        if extension in {".txt", ".md", ".markdown"}:
            return TextParser(filename, mime_type)
        try:
            return self._parsers[extension]
        except KeyError as exc:
            raise ValueError(f"Unsupported document extension: {extension or '(none)'}") from exc

    def availability(self) -> dict[str, bool]:
        return {
            extension: True
            for extension in sorted({".txt", ".md", *self._parsers})
        }


parser_registry = ParserRegistry()
parser_registry.register(".csv", CsvParser())
parser_registry.register(".pdf", PdfParser())
parser_registry.register(".docx", DocxParser())
parser_registry.register(".xlsx", XlsxParser())


def parser_for(filename: str, mime_type: str | None = None) -> Parser:
    return parser_registry.get(filename, mime_type)
